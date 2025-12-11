#!/usr/bin/env python3
"""
Ingest - Document Processing CLI for RAG Pipelines
Processes individual files or batches using various backends (marker-pdf, etc.) for high-quality extraction.
Supports markdown, json, html, and chunks output formats with optional LLM enhancement.
"""

import os
import json
import uuid
import time
import base64
import io
from pathlib import Path
from typing import List, Dict, Any, Tuple, Union, Optional
from concurrent.futures import ProcessPoolExecutor, as_completed

# Basic imports
from loguru import logger
from tqdm import tqdm

# Document cleaning
try:
    from document_cleaner import DocumentCleaner
    CLEANER_AVAILABLE = True
except ImportError:
    CLEANER_AVAILABLE = False
    logger.warning("document_cleaner not available, skipping text cleaning")

# Marker-pdf imports
try:
    from marker.models import create_model_dict
    from marker.converters.pdf import PdfConverter
    from marker.converters.table import TableConverter
    from marker.converters.ocr import OCRConverter
    from marker.config.parser import ConfigParser
    from marker.output import text_from_rendered
    from marker.schema import BlockTypes
    from PIL import Image
    MARKER_AVAILABLE = True
except ImportError:
    MARKER_AVAILABLE = False
    logger.warning("marker-pdf not available, falling back to basic processing")

# Fallback imports for when marker-pdf is not available
import pypdf
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

class MarkerDocumentProcessor:
    """
    Unified document processor that uses marker-pdf[full] for high-quality extraction.
    Supports multiple output formats, LLM enhancement, and specialized converters.
    Falls back to basic processing when marker-pdf is not available.
    """

    VERSION = "1.0.0"

    # Supported input formats (marker-pdf via PyMuPDF)
    SUPPORTED_EXTENSIONS = {
        # Document formats
        '.pdf', '.xps', '.oxps', '.epub', '.mobi', '.fb2', '.cbz', '.cbr',
        # Image formats (PyMuPDF can handle these as documents)
        '.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff', '.pnm', '.pgm',
        '.pbm', '.ppm', '.pam', '.jxr', '.jpx', '.jp2',
        # Text-based formats (opened as text documents)
        '.txt', '.md', '.rst', '.html', '.htm', '.xml', '.json', '.csv',
        # Office formats
        '.docx', '.xlsx', '.pptx'
    }

    SUPPORTED_OUTPUT_FORMATS = {'markdown', 'json', 'html', 'chunks'}
    SUPPORTED_CONVERTER_TYPES = {'pdf', 'table', 'ocr'}

    def __init__(
        self,
        output_dir: str = None,
        batch_mode: bool = False,
        use_marker: bool = True,
        output_format: str = "markdown",
        converter_type: str = "pdf",
        use_llm: bool = False,
        force_ocr: bool = False,
        page_range: Optional[str] = None,
        debug: bool = False,
        workers: int = 1,
        llm_service: Optional[str] = None,
        clean_output: bool = True,
        **marker_config
    ):
        """
        Initialize the document processor.

        Args:
            output_dir: Base output directory (default: ./processed)
            batch_mode: If True, creates unified batch outputs instead of individual folders
            use_marker: If True, uses marker-pdf when available
            output_format: Output format - 'markdown', 'json', 'html', or 'chunks'
            converter_type: Converter type - 'pdf', 'table', or 'ocr'
            use_llm: Use LLM to improve accuracy (requires LLM service configuration)
            force_ocr: Force OCR on all pages (also formats inline math)
            page_range: Specific pages to process (e.g., "0,5-10,20")
            debug: Enable debug mode (saves layout images and bbox JSON)
            workers: Number of parallel workers for batch processing
            llm_service: LLM service class path (e.g., 'marker.services.gemini.GoogleGeminiService')
            clean_output: Enable intelligent document cleaning (headers, footers, artifacts)
            **marker_config: Additional marker configuration options
        """
        # Validate parameters
        if output_format not in self.SUPPORTED_OUTPUT_FORMATS:
            raise ValueError(f"output_format must be one of {self.SUPPORTED_OUTPUT_FORMATS}")
        if converter_type not in self.SUPPORTED_CONVERTER_TYPES:
            raise ValueError(f"converter_type must be one of {self.SUPPORTED_CONVERTER_TYPES}")

        self.output_dir = Path(output_dir) if output_dir else Path("./processed")
        self.batch_mode = batch_mode
        self.use_marker = use_marker and MARKER_AVAILABLE
        self.output_format = output_format
        self.converter_type = converter_type
        self.workers = max(1, workers)
        self.clean_output = clean_output and CLEANER_AVAILABLE

        # Initialize document cleaner
        if self.clean_output:
            self.cleaner = DocumentCleaner()
            logger.info("Document cleaning enabled")
        else:
            self.cleaner = None

        self.processing_stats = {
            "total_files": 0,
            "processed": 0,
            "skipped": 0,
            "errors": [],
            "file_types": {},
            "processing_engine": "marker-pdf" if self.use_marker else "basic"
        }

        # Initialize marker-pdf models if available
        if self.use_marker:
            logger.info(f"Initializing marker-pdf with {converter_type} converter...")
            try:
                self.model_dict = create_model_dict()

                # Build configuration
                config = {
                    "output_format": output_format,
                    "force_ocr": force_ocr,
                    "use_llm": use_llm,
                    "debug": debug,
                    **marker_config
                }

                if page_range:
                    config["page_range"] = page_range

                self.config_parser = ConfigParser(config)

                # Initialize the appropriate converter
                converter_kwargs = {
                    "artifact_dict": self.model_dict,
                    "config": self.config_parser.generate_config_dict(),
                }

                # Add optional components if available
                try:
                    converter_kwargs["processor_list"] = self.config_parser.get_processors()
                    converter_kwargs["renderer"] = self.config_parser.get_renderer()
                    if use_llm:
                        converter_kwargs["llm_service"] = self.config_parser.get_llm_service()
                except Exception as e:
                    logger.warning(f"Could not initialize optional components: {e}")

                if converter_type == "pdf":
                    self.converter = PdfConverter(**converter_kwargs)
                elif converter_type == "table":
                    self.converter = TableConverter(**converter_kwargs)
                elif converter_type == "ocr":
                    self.converter = OCRConverter(**converter_kwargs)

                logger.success(f"Marker-pdf {converter_type} converter loaded successfully")
                logger.info(f"Output format: {output_format}, LLM: {use_llm}, Force OCR: {force_ocr}")

            except Exception as e:
                logger.error(f"Failed to load marker-pdf models: {e}")
                self.use_marker = False
                self.processing_stats["processing_engine"] = "basic"
        else:
            logger.info("Using basic processing (marker-pdf not available or disabled)")
    
    def process(self, inputs: Union[str, List[str], Path]) -> Dict[str, Dict[str, str]]:
        """
        Universal processing method that handles different input types.

        Args:
            inputs: Can be:
                - str: single file path or directory path
                - list: list of file paths
                - Path: pathlib Path object

        Returns:
            Dictionary mapping input files to their output paths:
            {
                "file1.pdf": {"output": "path", "metadata": "path", "images": ["path1", ...]},
                "file2.xlsx": {"output": "path", "metadata": "path"}
            }
        """
        logger.info(f"Starting document processing with engine={self.processing_stats['processing_engine']}")
        logger.info(f"Output format: {self.output_format}, Converter: {self.converter_type}, Workers: {self.workers}")

        # Normalize inputs to list of file paths
        file_paths = self._normalize_inputs(inputs)

        if not file_paths:
            raise ValueError("No valid files found to process")

        logger.info(f"Found {len(file_paths)} files to process")

        # Prepare output directories
        if self.batch_mode:
            self._prepare_batch_output()
        else:
            self.output_dir.mkdir(parents=True, exist_ok=True)

        # Process files (with multi-worker support)
        if self.workers > 1 and len(file_paths) > 1:
            results = self._process_files_parallel(file_paths)
        else:
            results = self._process_files_sequential(file_paths)

        # Generate batch summary if in batch mode
        if self.batch_mode:
            self._generate_batch_summary()

        logger.success(f"Processing complete: {self.processing_stats['processed']}/{len(file_paths)} files processed")
        if self.processing_stats["errors"]:
            logger.warning(f"Encountered {len(self.processing_stats['errors'])} errors")
        return results

    def _process_files_sequential(self, file_paths: List[Path]) -> Dict[str, Dict[str, str]]:
        """Process files sequentially."""
        results = {}
        for file_path in tqdm(file_paths, desc="Processing files"):
            try:
                result = self._process_single_file(file_path)
                results[str(file_path)] = result
                self.processing_stats["processed"] += 1
            except Exception as e:
                error_msg = f"Failed to process {file_path.name}: {e}"
                logger.error(error_msg)
                self.processing_stats["errors"].append(error_msg)
        return results

    def _process_files_parallel(self, file_paths: List[Path]) -> Dict[str, Dict[str, str]]:
        """Process files in parallel using multiple workers."""
        logger.info(f"Processing {len(file_paths)} files with {self.workers} workers")
        results = {}

        with ProcessPoolExecutor(max_workers=self.workers) as executor:
            # Submit all tasks
            future_to_path = {
                executor.submit(self._process_single_file, file_path): file_path
                for file_path in file_paths
            }

            # Process completed tasks with progress bar
            with tqdm(total=len(file_paths), desc="Processing files") as pbar:
                for future in as_completed(future_to_path):
                    file_path = future_to_path[future]
                    try:
                        result = future.result()
                        results[str(file_path)] = result
                        self.processing_stats["processed"] += 1
                    except Exception as e:
                        error_msg = f"Failed to process {file_path.name}: {e}"
                        logger.error(error_msg)
                        self.processing_stats["errors"].append(error_msg)
                    finally:
                        pbar.update(1)

        return results
    
    def _normalize_inputs(self, inputs: Union[str, List[str], Path]) -> List[Path]:
        """Convert various input types to a list of file paths."""
        file_paths = []
        
        if isinstance(inputs, str):
            input_path = Path(inputs)
            if input_path.is_dir():
                # Directory: find all supported files
                file_paths = self._find_files_in_directory(input_path)
            elif input_path.exists():
                # Single file
                file_paths = [input_path]
            else:
                raise FileNotFoundError(f"Input path does not exist: {inputs}")
        
        elif isinstance(inputs, Path):
            if inputs.is_dir():
                file_paths = self._find_files_in_directory(inputs)
            elif inputs.exists():
                file_paths = [inputs]
            else:
                raise FileNotFoundError(f"Input path does not exist: {inputs}")
        
        elif isinstance(inputs, list):
            # List of files
            for item in inputs:
                path = Path(item)
                if path.exists():
                    file_paths.append(path)
                else:
                    logger.warning(f"File not found, skipping: {item}")
        
        else:
            raise ValueError(f"Unsupported input type: {type(inputs)}")
        
        # Filter by supported extensions
        supported_files = []
        for path in file_paths:
            if path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                supported_files.append(path)
                ext = path.suffix.lower()
                self.processing_stats["file_types"][ext] = self.processing_stats["file_types"].get(ext, 0) + 1
            else:
                logger.warning(f"Unsupported file type, skipping: {path.name}")
        
        self.processing_stats["total_files"] = len(supported_files)
        
        # Log file type breakdown
        for ext, count in self.processing_stats["file_types"].items():
            if count > 0:
                logger.info(f"  {ext}: {count} files")
        
        return supported_files
    
    def _find_files_in_directory(self, directory: Path) -> List[Path]:
        """Recursively find all supported files in directory."""
        all_files = []
        for ext in self.SUPPORTED_EXTENSIONS:
            files = list(directory.rglob(f"*{ext}"))
            all_files.extend(files)
        return all_files
    
    def _prepare_batch_output(self):
        """Prepare output directories for batch mode."""
        self.output_dir.mkdir(parents=True, exist_ok=True)
        (self.output_dir / self.output_format).mkdir(exist_ok=True)
        (self.output_dir / "metadata").mkdir(exist_ok=True)
        (self.output_dir / "images").mkdir(exist_ok=True)
    
    def _process_single_file(self, file_path: Path) -> Dict[str, str]:
        """Process a single file and return output paths."""
        # Determine output extension based on format
        format_extensions = {
            "markdown": ".md",
            "json": ".json",
            "html": ".html",
            "chunks": ".json"
        }
        output_ext = format_extensions.get(self.output_format, ".txt")

        # Determine output paths
        if self.batch_mode:
            base_name = file_path.stem
            output_path = self.output_dir / self.output_format / f"{base_name}{output_ext}"
            metadata_path = self.output_dir / "metadata" / f"{base_name}_metadata.json"
            images_dir = self.output_dir / "images"
        else:
            # Individual folder per file
            output_folder = self.output_dir / f"{file_path.stem}_processed"
            output_folder.mkdir(parents=True, exist_ok=True)
            output_path = output_folder / f"content{output_ext}"
            metadata_path = output_folder / "metadata.json"
            images_dir = output_folder / "images"

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        metadata_path.parent.mkdir(parents=True, exist_ok=True)

        # Process content based on availability
        if self.use_marker and file_path.suffix.lower() in ['.pdf', '.xps', '.oxps', '.epub', '.mobi', '.fb2', '.cbz', '.cbr', '.png', '.jpg', '.jpeg', '.pptx', '.docx', '.xlsx', '.html', '.htm']:
            # Use marker-pdf for supported formats
            content, metadata, image_paths = self._process_with_marker(file_path)
        else:
            # Use basic processing for other formats or fallback
            content, metadata = self._process_basic(file_path)
            image_paths = []

        # Add base file metadata
        base_metadata = self._extract_metadata(file_path)
        metadata.update(base_metadata)

        # Clean content if enabled
        if self.clean_output and self.cleaner:
            if self.output_format in ["markdown", "html"]:
                content = self.cleaner.clean_markdown(content)
            # For JSON/chunks, content is already structured

        # Write outputs
        if self.output_format in ["markdown", "html"]:
            output_path.write_text(content, encoding='utf-8')
        else:  # json or chunks
            output_path.write_text(json.dumps(content, indent=2, ensure_ascii=False), encoding='utf-8')

        metadata_path.write_text(json.dumps(metadata, indent=2, ensure_ascii=False), encoding='utf-8')

        result = {
            "output": str(output_path),
            "metadata": str(metadata_path)
        }

        if image_paths:
            result["images"] = image_paths

        return result
    
    def _process_with_marker(self, file_path: Path) -> Tuple[Any, Dict, List]:
        """
        Process file using marker-pdf for high-quality extraction.

        Returns:
            Tuple of (content, metadata, images)
            - content: text string for markdown/html, structured data for json/chunks
            - metadata: extracted metadata dict
            - images: list of saved image paths
        """
        logger.info(f"Processing {file_path.name} with marker-pdf ({self.output_format})")

        # Run the converter
        rendered = self.converter(str(file_path))

        # Extract content based on output format
        if self.output_format in ["markdown", "html"]:
            # For markdown/html, use text_from_rendered
            text, marker_metadata, images_dict = text_from_rendered(rendered)
            content = text

        elif self.output_format == "json":
            # For JSON, return the structured block tree
            content = {
                "pages": rendered.children if hasattr(rendered, 'children') else [],
                "block_type": rendered.block_type if hasattr(rendered, 'block_type') else None
            }
            marker_metadata = rendered.metadata if hasattr(rendered, 'metadata') else {}
            images_dict = rendered.images if hasattr(rendered, 'images') else {}

        elif self.output_format == "chunks":
            # For chunks, return flattened structure (optimized for RAG)
            if hasattr(rendered, 'chunks'):
                content = rendered.chunks
            else:
                # Fallback to extracting top-level blocks
                content = []
                if hasattr(rendered, 'children'):
                    for page in rendered.children:
                        if hasattr(page, 'children'):
                            content.extend(page.children)
            marker_metadata = rendered.metadata if hasattr(rendered, 'metadata') else {}
            images_dict = rendered.images if hasattr(rendered, 'images') else {}

        else:
            raise ValueError(f"Unsupported output format: {self.output_format}")

        # Save images if present
        saved_image_paths = []
        if images_dict:
            saved_image_paths = self._save_images(images_dict, file_path.stem)

        # Build comprehensive metadata
        file_metadata = self._extract_file_metadata(file_path)
        combined_metadata = {
            **file_metadata,
            "marker_metadata": marker_metadata,
            "processing_engine": "marker-pdf",
            "converter_type": self.converter_type,
            "output_format": self.output_format
        }

        return content, combined_metadata, saved_image_paths
    
    def _process_basic(self, file_path: Path) -> Tuple[str, Dict]:
        """
        Process file using basic libraries (fallback).

        Returns:
            Tuple of (content, metadata)
        """
        logger.info(f"Processing {file_path.name} with basic processor")

        file_ext = file_path.suffix.lower()

        if file_ext == '.pdf':
            return self._process_pdf_basic(file_path)
        elif file_ext == '.docx' and DOCX_AVAILABLE:
            return self._process_docx_basic(file_path)
        elif file_ext == '.xlsx' and XLSX_AVAILABLE:
            return self._process_xlsx_basic(file_path)
        elif file_ext in ['.txt', '.md', '.rst', '.html', '.htm', '.xml', '.json', '.csv']:
            return self._process_text_basic(file_path)
        else:
            # Generic fallback
            return self._process_generic_basic(file_path)
    
    def _save_images(self, images_dict: Dict, file_stem: str) -> List[str]:
        """
        Save images from marker-pdf output.

        Args:
            images_dict: Dictionary of images from marker (block_id -> image data)
            file_stem: Base filename without extension

        Returns:
            List of saved image file paths
        """
        saved_paths = []

        if self.batch_mode:
            images_dir = self.output_dir / "images"
        else:
            images_dir = self.output_dir / f"{file_stem}_processed" / "images"

        images_dir.mkdir(parents=True, exist_ok=True)

        for block_id, image_data in images_dict.items():
            try:
                # Create safe filename from block_id
                safe_block_id = block_id.replace('/', '_').replace('\\', '_')
                safe_name = f"{file_stem}_{safe_block_id}.png"
                save_path = images_dir / safe_name

                # marker-pdf provides PIL Image objects or base64 strings
                if hasattr(image_data, 'save'):
                    # PIL Image object
                    image_data.save(save_path)
                    saved_paths.append(str(save_path))

                elif isinstance(image_data, str):
                    # Base64 string
                    image_bytes = base64.b64decode(image_data)
                    image = Image.open(io.BytesIO(image_bytes))
                    image.save(save_path)
                    saved_paths.append(str(save_path))

                elif isinstance(image_data, bytes):
                    # Raw bytes
                    image = Image.open(io.BytesIO(image_data))
                    image.save(save_path)
                    saved_paths.append(str(save_path))

            except Exception as e:
                logger.warning(f"Could not save image {block_id}: {e}")

        if saved_paths:
            logger.info(f"Saved {len(saved_paths)} images to {images_dir}")

        return saved_paths
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract comprehensive metadata from file."""
        metadata = {
            "file_info": {
                "name": file_path.name,
                "path": str(file_path),
                "size": file_path.stat().st_size,
                "type": file_path.suffix.lower(),
                "modified": file_path.stat().st_mtime,
                "created": file_path.stat().st_ctime
            },
            "processing": {
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime()),
                "processor": "ingest",
                "version": self.VERSION,
                "batch_mode": self.batch_mode,
                "engine": self.processing_stats["processing_engine"],
                "output_format": self.output_format,
                "converter_type": self.converter_type if self.use_marker else None
            }
        }

        # Add document-specific metadata
        doc_metadata = self._extract_file_metadata(file_path)
        if doc_metadata:
            metadata["document"] = doc_metadata

        return metadata
    
    def _extract_file_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document-specific metadata."""
        file_ext = file_path.suffix.lower()
        metadata = {}
        
        try:
            if file_ext == '.pdf' and not self.use_marker:
                # Basic PDF metadata extraction
                with open(file_path, "rb") as f:
                    pdf_reader = pypdf.PdfReader(f)
                    if pdf_reader.metadata:
                        metadata = {
                            "author": pdf_reader.metadata.get("/Author", ""),
                            "date": pdf_reader.metadata.get("/CreationDate", ""),
                            "title": pdf_reader.metadata.get("/Title", ""),
                            "creator": pdf_reader.metadata.get("/Creator", ""),
                            "producer": pdf_reader.metadata.get("/Producer", ""),
                            "page_count": len(pdf_reader.pages)
                        }
            
            elif file_ext == '.docx' and DOCX_AVAILABLE:
                doc = DocxDocument(file_path)
                metadata = {
                    "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "table_count": len(doc.tables),
                    "page_count": 1
                }
                
                if doc.core_properties:
                    core_props = doc.core_properties
                    metadata.update({
                        "title": getattr(core_props, 'title', '') or '',
                        "author": getattr(core_props, 'author', '') or '',
                        "created": str(getattr(core_props, 'created', '')) or '',
                        "modified": str(getattr(core_props, 'modified', '')) or ''
                    })
            
            elif file_ext == '.xlsx' and XLSX_AVAILABLE:
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                metadata = {
                    "sheet_count": len(workbook.sheetnames),
                    "sheet_names": workbook.sheetnames,
                    "page_count": 1
                }
                workbook.close()
                
        except Exception as e:
            logger.warning(f"Could not extract metadata for {file_path.name}: {e}")
        
        return metadata
    
    # Basic processing methods (fallback when marker-pdf is not available)
    def _process_pdf_basic(self, file_path: Path) -> Tuple[str, Dict]:
        """Basic PDF processing using pypdf."""
        pdf_reader = pypdf.PdfReader(str(file_path))

        content = f"# {file_path.name}\n\n"

        for page_num, page in enumerate(pdf_reader.pages, 1):
            page_text = page.extract_text()
            content += f"## Page {page_num}\n\n"
            content += page_text + "\n\n"

        metadata = {
            "file_type": ".pdf",
            "page_count": len(pdf_reader.pages),
            "processing_engine": "basic"
        }

        return content, metadata
    
    def _process_docx_basic(self, file_path: Path) -> Tuple[str, Dict]:
        """Basic DOCX processing using python-docx."""
        doc = DocxDocument(file_path)

        content = f"# {file_path.name}\n\n"
        content += "## Document Content\n\n"

        for para in doc.paragraphs:
            if para.text.strip():
                content += para.text + "\n\n"

        if doc.tables:
            content += "### Tables\n\n"
            for table_idx, table in enumerate(doc.tables, 1):
                content += f"**Table {table_idx}:**\n\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    content += row_text + "\n"
                content += "\n"

        metadata = {
            "file_type": ".docx",
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
            "table_count": len(doc.tables),
            "processing_engine": "basic"
        }

        return content, metadata
    
    def _process_xlsx_basic(self, file_path: Path) -> Tuple[str, Dict]:
        """Basic XLSX processing using openpyxl."""
        workbook = openpyxl.load_workbook(file_path, read_only=True)

        content = f"# {file_path.name}\n\n"
        content += "## Excel Content\n\n"

        sheet_count = len(workbook.sheetnames)

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            content += f"### Sheet: {sheet_name}\n\n"

            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    content += row_text + "\n"

            content += "\n"

        workbook.close()

        metadata = {
            "file_type": ".xlsx",
            "sheet_count": sheet_count,
            "processing_engine": "basic"
        }

        return content, metadata
    
    def _process_text_basic(self, file_path: Path) -> Tuple[str, Dict]:
        """Basic text file processing."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text_content = f.read()

        content = f"# {file_path.name}\n\n"
        content += "## Content\n\n"
        content += "```\n" + text_content + "\n```\n\n"

        metadata = {
            "file_type": file_path.suffix.lower(),
            "character_count": len(text_content),
            "processing_engine": "basic"
        }

        return content, metadata

    def _process_generic_basic(self, file_path: Path) -> Tuple[str, Dict]:
        """Generic fallback for unsupported file types."""
        content = f"# {file_path.name}\n\n"
        content += f"## File Information\n\n"
        content += f"- **Type**: {file_path.suffix}\n"
        content += f"- **Size**: {file_path.stat().st_size} bytes\n"
        content += f"- **Processing**: Basic fallback (binary file)\n\n"

        metadata = {
            "file_type": file_path.suffix.lower(),
            "processing_engine": "basic"
        }

        return content, metadata
    
    def _generate_batch_summary(self):
        """Generate batch processing summary."""
        summary = {
            "batch_info": {
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S", time.gmtime()),
                "output_directory": str(self.output_dir),
                "batch_mode": True,
                "processing_engine": self.processing_stats["processing_engine"]
            },
            "statistics": self.processing_stats,
            "configuration": {
                "supported_formats": list(self.SUPPORTED_EXTENSIONS),
                "processor_version": self.VERSION,
                "marker_available": MARKER_AVAILABLE,
                "use_marker": self.use_marker,
                "output_format": self.output_format,
                "converter_type": self.converter_type,
                "workers": self.workers
            }
        }

        summary_path = self.output_dir / "batch_summary.json"
        summary_path.write_text(json.dumps(summary, indent=2, ensure_ascii=False))
        logger.info(f"Batch summary saved to: {summary_path}")

def main():
    """Command line interface for the unified marker processor."""
    import argparse

    parser = argparse.ArgumentParser(
        description="Ingest - Document Processing CLI for RAG Pipelines v1.0",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Process a single PDF with default settings
  %(prog)s document.pdf

  # Process directory with LLM enhancement and chunks output
  %(prog)s ./documents --output-format chunks --use-llm --batch-mode

  # Extract tables only with multi-worker processing
  %(prog)s ./pdfs --converter-type table --workers 4 --batch-mode

  # Process with forced OCR and debug mode
  %(prog)s file.pdf --force-ocr --debug --output-format markdown
        """
    )

    # Input/Output arguments
    parser.add_argument("inputs", nargs="+", help="Files or directories to process")
    parser.add_argument("--output-dir", help="Output directory (default: ./processed)")
    parser.add_argument("--batch-mode", action="store_true",
                       help="Create unified batch outputs instead of individual folders")

    # Output format options
    parser.add_argument("--output-format", choices=["markdown", "json", "html", "chunks"],
                       default="markdown", help="Output format (default: markdown)")
    parser.add_argument("--converter-type", choices=["pdf", "table", "ocr"],
                       default="pdf", help="Converter type (default: pdf)")

    # Marker-pdf options
    parser.add_argument("--no-marker", action="store_true",
                       help="Disable marker-pdf and use basic processing")
    parser.add_argument("--use-llm", action="store_true",
                       help="Use LLM to improve accuracy (requires LLM service configuration)")
    parser.add_argument("--force-ocr", action="store_true",
                       help="Force OCR on all pages (also formats inline math)")
    parser.add_argument("--page-range", type=str,
                       help="Specific pages to process (e.g., '0,5-10,20')")
    parser.add_argument("--debug", action="store_true",
                       help="Enable debug mode (saves layout images and bbox JSON)")

    # Performance options
    parser.add_argument("--workers", type=int, default=1,
                       help="Number of parallel workers for batch processing (default: 1)")

    # LLM service options
    parser.add_argument("--llm-service", type=str,
                       help="LLM service class path (e.g., 'marker.services.gemini.GoogleGeminiService')")
    parser.add_argument("--gemini-api-key", type=str,
                       help="Gemini API key for LLM mode")
    parser.add_argument("--claude-api-key", type=str,
                       help="Claude API key for LLM mode")

    # Additional marker config
    parser.add_argument("--disable-image-extraction", action="store_true",
                       help="Don't extract images from documents")
    parser.add_argument("--strip-existing-ocr", action="store_true",
                       help="Remove existing OCR text and re-OCR")
    parser.add_argument("--paginate-output", action="store_true",
                       help="Add page breaks to output")

    # Document cleaning
    parser.add_argument("--no-clean", action="store_true",
                       help="Disable intelligent document cleaning (headers/footers/artifacts)")

    args = parser.parse_args()

    try:
        # Build marker config from CLI args
        marker_config = {}
        if args.disable_image_extraction:
            marker_config["disable_image_extraction"] = True
        if args.strip_existing_ocr:
            marker_config["strip_existing_ocr"] = True
        if args.paginate_output:
            marker_config["paginate_output"] = True
        if args.gemini_api_key:
            marker_config["gemini_api_key"] = args.gemini_api_key
        if args.claude_api_key:
            marker_config["claude_api_key"] = args.claude_api_key

        # Initialize processor
        processor = MarkerDocumentProcessor(
            output_dir=args.output_dir,
            batch_mode=args.batch_mode,
            use_marker=not args.no_marker,
            output_format=args.output_format,
            converter_type=args.converter_type,
            use_llm=args.use_llm,
            force_ocr=args.force_ocr,
            page_range=args.page_range,
            debug=args.debug,
            workers=args.workers,
            llm_service=args.llm_service,
            clean_output=not args.no_clean,
            **marker_config
        )

        # Process inputs
        if len(args.inputs) == 1:
            result = processor.process(args.inputs[0])
        else:
            result = processor.process(args.inputs)

        # Display results
        print(f"\nProcessing complete!")
        print(f"Engine: {processor.processing_stats['processing_engine']}")
        print(f"Output format: {args.output_format}")
        print(f"Processed {len(result)} files:")
        for input_file, outputs in result.items():
            print(f"  {Path(input_file).name}")
            print(f"    Output: {outputs['output']}")
            print(f"    Metadata: {outputs['metadata']}")
            if 'images' in outputs:
                print(f"    Images: {len(outputs['images'])} files")

        if processor.processing_stats["errors"]:
            print(f"\nEncountered {len(processor.processing_stats['errors'])} errors")
            for error in processor.processing_stats["errors"][:5]:  # Show first 5 errors
                print(f"  - {error}")

        return 0

    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    exit(main())
